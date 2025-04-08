import unicodedata
from pathlib import Path
from typing import List, Optional
import pandas as pd
from docx import Document as DocxDocument

def docx_reader(file_path: Path, extra_info: Optional[dict] = None) -> List[dict]:
    """
    Load data from a .docx file, combining paragraphs and tables. Each paragraph is paired with a relevant table,
    and the resultant data includes metadata about the document.

    Args:
        file_path (Path): Path to the .docx file.
        extra_info (dict, optional): Additional metadata to be added to each document.

    Returns:
        List[dict]: A list of dictionaries representing the extracted documents (paragraphs and tables).
    """
    file_path = Path(file_path).resolve()
    file_name = file_path.name
    file_metadata = {"file_name": file_name, "file_path": str(file_path), "type": "docx"}
    
    # Load the docx document
    doc = DocxDocument(str(file_path))
    
    # Extract paragraphs and normalize text
    all_paragraphs = [unicodedata.normalize("NFKC", p.text) for p in doc.paragraphs if p.text.strip()]
    
    # Extract tables as structured data
    tables = []
    for table in doc.tables:
        table_data = _load_single_table(table)
        df = pd.DataFrame(table_data[1:], columns=table_data[0]) if table_data else pd.DataFrame()
        tables.append(df)
    
    extra_info = extra_info or {}
    
    documents = []
    paragraph_idx = 0  # Track paragraph index to associate with tables

    # Combine paragraphs and their relevant tables
    for paragraph in all_paragraphs:
        table = tables[paragraph_idx] if paragraph_idx < len(tables) else None

        # Prepare the document entry
        document_entry = {
            "text": paragraph,
            "metadata": {"type": "paragraph",
                          "page_label": paragraph_idx + 1,
                          "table_data": table.to_csv(index=False).strip() if table is not None and not table.empty else None,
                          **file_metadata,
                          **extra_info}
        }

        # Attach table if available
        

        documents.append(document_entry)

        # Move to the next paragraph-table pair
        paragraph_idx += 1

    return documents

def _load_single_table(table) -> List[List[str]]:
    """
    Extract content from tables in a .docx file. Return a list of columns where each column is a list of strings.

    Args:
        table: A `table` element from the docx document.

    Returns:
        List[List[str]]: A list of columns, where each column is a list of strings.
    """
    if not table.rows or not table.columns:
        return []
    
    return [[cell.text.strip() for cell in row.cells] for row in table.rows]

# Example usage:
# file_path = Path("example_document.docx")
# documents = docx_reader(file_path)
# print(documents)
